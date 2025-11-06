import io
import os
import re
from fastapi import FastAPI, Form, Request, UploadFile
from fastapi.responses import HTMLResponse, StreamingResponse, RedirectResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from docx import Document
from dotenv import load_dotenv

import google.generativeai as genai

load_dotenv()
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))

app = FastAPI()
templates = Jinja2Templates(directory="templates")

# In-memory state
SESSION = {}

# Debug endpoint to check session state (remove in production)
@app.get("/debug/session")
async def debug_session():
    return {"session_keys": list(SESSION.keys()), "session_data": SESSION}


def extract_text_from_docx(file: UploadFile) -> tuple[str, Document]:
    """Extract text and return both text and the document object"""
    # Reset file pointer to beginning
    file.file.seek(0)
    doc = Document(file.file)
    full_text = []
    for p in doc.paragraphs:
        full_text.append(p.text)
    # Also check tables for text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
    return "\n".join(full_text), doc


def extract_placeholders(text: str):
    patterns = [
        r"\[[^\]]+\]",
        r"\$?\[[_.\s]*\]",
        r"\{[^\}]+\}",
        r"<[^>]+>",
    ]
    placeholders = set()
    for pat in patterns:
        placeholders.update(re.findall(pat, text))
    return list(placeholders)


def identify_placeholders_with_llm(text: str, placeholders: list):
    """
    Use Gemini to identify what each placeholder means in context.
    """
    snippets = {}
    for ph in placeholders:
        for match in re.finditer(re.escape(ph), text):
            start = max(0, match.start() - 100)
            end = min(len(text), match.end() + 100)
            snippets[ph] = text[start:end]

    prompt = f"""
You are a legal document analyzer. For each placeholder, describe what the user should fill in.

Return **valid JSON only** in this format whaich would be passed to `json.loads()`:
[
  {{"placeholder": "...", "label": "...", "question": "..."}}
]

Snippets:
{snippets}
"""

    try:
        model = genai.GenerativeModel("gemini-2.0-flash")
        response = model.generate_content(prompt)
        
        import json
        data = json.loads(response.text.removeprefix("```json").removesuffix("```").strip())
    except Exception as e:
        print(f"Error with Gemini API: {e}")
        # Fallback if API call or parsing fails
        data = [
            {"placeholder": ph, "label": ph, "question": f"What should I fill for {ph}?"}
            for ph in placeholders
        ]

    return data


def replace_placeholders_in_document(doc: Document, answers: dict):
    """
    Replace placeholders in document while preserving formatting.
    This function modifies paragraphs and table cells in place.
    """
    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        for placeholder, answer in answers.items():
            if placeholder in paragraph.text:
                # Replace text while preserving formatting
                replace_text_in_paragraph(paragraph, placeholder, answer)
    
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for placeholder, answer in answers.items():
                        if placeholder in paragraph.text:
                            replace_text_in_paragraph(paragraph, placeholder, answer)


def replace_text_in_paragraph(paragraph, old_text, new_text):
    """
    Replace text in a paragraph while preserving formatting.
    This handles cases where the placeholder might span multiple runs.
    """
    if old_text not in paragraph.text:
        return
    
    # Get all runs in the paragraph
    runs = paragraph.runs
    
    # Build full text and track run boundaries
    full_text = ""
    run_boundaries = []
    for run in runs:
        start = len(full_text)
        full_text += run.text
        end = len(full_text)
        run_boundaries.append((start, end, run))
    
    # Find placeholder positions
    start_pos = full_text.find(old_text)
    if start_pos == -1:
        return
    
    end_pos = start_pos + len(old_text)
    
    # Clear all run texts first
    for _, _, run in run_boundaries:
        run.text = ""
    
    # Rebuild text with replacement
    new_full_text = full_text.replace(old_text, new_text)
    
    # Put the new text in the first run to preserve some formatting
    if runs:
        runs[0].text = new_full_text


@app.get("/", response_class=HTMLResponse)
async def index(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})


@app.post("/upload", response_class=HTMLResponse)
async def upload(request: Request, file: UploadFile):
    text, doc = extract_text_from_docx(file)
    placeholders = extract_placeholders(text)
    details = identify_placeholders_with_llm(text, placeholders)

    SESSION["text"] = text
    SESSION["doc"] = doc  # Store the original document object
    SESSION["placeholders"] = details
    SESSION["answers"] = {}
    SESSION["file_name"] = file.filename

    # Return analysis result instead of going directly to chat
    filename_display = file.filename[:30] + "..." if len(file.filename) > 30 else file.filename
    
    return templates.TemplateResponse(
        "analysis.html",
        {
            "request": request,
            "total_placeholders": len(details),
            "filename_display": filename_display,
        },
    )


@app.get("/chat", response_class=HTMLResponse)
async def start_chat(request: Request):
    if "placeholders" not in SESSION:
        return RedirectResponse(url="/", status_code=302)
    
    details = SESSION["placeholders"]
    first_q = details[0]["question"] if details else "No placeholders found."
    total_questions = len(details)
    progress_percentage = round((0 + 1) / total_questions * 100, 1) if total_questions > 0 else 0
    
    return templates.TemplateResponse(
        "chat.html",
        {
            "request": request, 
            "index": 0, 
            "question": first_q, 
            "total_questions": total_questions,
            "progress_percentage": progress_percentage
        },
    )


@app.post("/fill", response_class=HTMLResponse)
async def fill(
    request: Request,
    answer: str = Form(...),
    index: int = Form(...)
):
    details = SESSION["placeholders"]
    SESSION["answers"][details[index]["placeholder"]] = answer

    index += 1
    if index < len(details):
        next_q = details[index]["question"]
        total_questions = len(details)
        
        # Return just the chat content with confirmation - wrapped in space-y-6 div
        return HTMLResponse(
            content=f"""<div class="space-y-6">
            <!-- Previous Answer Confirmation -->
            <div class="flex items-start space-x-4 mb-6 chat-bubble">
              <div class="flex-shrink-0 w-10 h-10 bg-gradient-to-r from-green-500 to-teal-600 rounded-full flex items-center justify-center">
                <i class="fas fa-check text-white"></i>
              </div>
              <div class="flex-1">
                <div class="bg-green-100 border border-green-200 rounded-2xl rounded-bl-sm p-4 shadow-sm">
                  <p class="text-green-800">✅ Got it! Your answer has been saved.</p>
                </div>
              </div>
            </div>
            
            <!-- Next Question -->
            <div class="flex items-start space-x-4">
              <div class="flex-shrink-0 w-10 h-10 bg-gradient-to-r from-blue-500 to-purple-600 rounded-full flex items-center justify-center">
                <i class="fas fa-robot text-white"></i>
              </div>
              <div class="flex-1">
                <div class="bg-slate-100 rounded-2xl rounded-tl-sm p-6 border border-slate-200 shadow-lg">
                  <div class="mb-4">
                    <div class="flex items-center justify-between mb-3">
                      <span class="text-slate-600 text-sm font-medium">AI Assistant</span>
                      <span class="text-blue-600 text-sm font-semibold">{index + 1}/{len(details)}</span>
                    </div>
                    <p class="text-slate-800 text-lg leading-relaxed font-medium">{next_q}</p>
                  </div>
                  <div class="flex items-center text-slate-500 text-sm">
                    <i class="fas fa-lightbulb mr-2"></i>
                    <span>Provide the most accurate information possible</span>
                  </div>
                </div>
              </div>
            </div>

            <!-- User Input Form -->
            <form hx-post="/fill" hx-target="#chat-box" hx-swap="innerHTML" class="flex items-end space-x-4 mt-6">
              <div class="flex-shrink-0 w-10 h-10 bg-gradient-to-r from-green-500 to-teal-600 rounded-full flex items-center justify-center">
                <i class="fas fa-user text-white"></i>
              </div>
              <div class="flex-1">
                <div class="bg-white rounded-2xl rounded-bl-sm p-4 border border-slate-200 shadow-lg">
                  <textarea 
                    name="answer" 
                    placeholder="Type your answer here..." 
                    class="w-full bg-transparent text-slate-800 placeholder-slate-500 resize-none focus:outline-none text-lg answer-textarea"
                    rows="3"
                    required
                  ></textarea>
                  <input type="hidden" name="index" value="{index}" />
                  <div class="flex items-center justify-between mt-4">
                    <div class="flex items-center text-slate-500 text-sm">
                      <i class="fas fa-info-circle mr-2"></i>
                      <span>Press Enter to send, Shift+Enter for new line</span>
                    </div>
                    <button type="submit" class="btn-primary font-semibold py-3 px-6 rounded-xl flex items-center space-x-2">
                      <span>Send</span>
                      <i class="fas fa-paper-plane"></i>
                    </button>
                  </div>
                </div>
              </div>
            </form>
            
            <!-- Simple Progress Info inside chat-box so it gets updated -->
            <div class="glass-effect rounded-2xl p-4 text-center mt-6">
              <div class="flex items-center justify-center space-x-2 text-slate-600">
                <i class="fas fa-clipboard-list text-blue-500"></i>
                <span class="text-sm">Question {index + 1} of {total_questions}</span>
                <span class="text-xs text-slate-400">•</span>
                <span class="text-sm">Keep going! You're doing great.</span>
              </div>
            </div>
            
            <script>
            // Auto-focus on textarea
            document.addEventListener('DOMContentLoaded', function() {{
              const textarea = document.querySelector('textarea[name="answer"]');
              if (textarea) textarea.focus();
            }});
            </script>
            </div>""",
            headers={"HX-Trigger": "updateProgress"}
        )
    else:
        # All questions completed - render completion page
        return templates.TemplateResponse(
            "complete.html",
            {
                "request": request,
                "filename": SESSION.get("file_name", "Unknown"),
                "total_fields": len(SESSION.get("placeholders", [])),
                "processing_time": "Less than 1 minute",
            },
        )


@app.get("/download")
async def download():
    # Check if session has required data
    if "text" not in SESSION:
        return HTMLResponse(
            content="<h2>Error: No document data found</h2><p>Please upload and fill a document first.</p><a href='/'>Go back to upload</a>",
            status_code=400
        )
    
    if "answers" not in SESSION:
        return HTMLResponse(
            content="<h2>Error: No answers found</h2><p>Please complete the form filling process first.</p><a href='/'>Go back to upload</a>",
            status_code=400
        )
    
    if "doc" not in SESSION:
        return HTMLResponse(
            content="<h2>Error: Original document not found</h2><p>Please upload a document first.</p><a href='/'>Go back to upload</a>",
            status_code=400
        )
    
    try:
        # Get the original document and answers
        original_doc = SESSION["doc"]
        answers = SESSION["answers"]
        file_name = SESSION.get("file_name", "document.docx")
        
        # Create a deep copy of the document by saving and reloading it
        # This ensures we don't modify the original stored document
        temp_buffer = io.BytesIO()
        original_doc.save(temp_buffer)
        temp_buffer.seek(0)
        
        # Create a new document from the saved buffer
        doc = Document(temp_buffer)
        
        # Replace placeholders in the document while preserving formatting
        replace_placeholders_in_document(doc, answers)
        
        # Save the final document to BytesIO
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        
        return StreamingResponse(
            io.BytesIO(output.getvalue()),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename=filled_{file_name}",
                "Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            }
        )
        
    except Exception as e:
        return HTMLResponse(
            content=f"<h2>Error creating document</h2><p>An error occurred: {str(e)}</p><a href='/'>Go back to upload</a>",
            status_code=500
        )


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
